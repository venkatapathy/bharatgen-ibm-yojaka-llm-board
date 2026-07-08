"""Export helpers for generated questions."""

from io import BytesIO

from docx import Document


def build_docx(run):
    document = Document()
    document.add_heading(run.name, level=1)
    document.add_paragraph(f"Topic: {run.topic}")
    for index, question in enumerate(run.questions.all(), start=1):
        document.add_paragraph(
            f"{index}. [{question.question_type}] ({question.marks} marks) {question.question_text}"
        )
        if question.options:
            for option in question.options:
                document.add_paragraph(f"- {option}", style="List Bullet")
        if question.reference_answer:
            document.add_paragraph(f"Answer: {question.reference_answer}")
    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream
